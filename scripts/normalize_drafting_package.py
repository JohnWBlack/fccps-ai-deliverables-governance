#!/usr/bin/env python3
"""Normalize drafting DOCX contributions and assemble an ordered policy package.

Behavior:
- Uses the provided outline DOCX as the base so Section 0 preface stays intact.
- Inserts a divider page for each section A-O.
- Places the section chair preface immediately after the divider page.
- Places contribution drafts after the chair preface.
- Appends a workstream preface section and workstream sections (WS-XXX) from D-XXX-1 docs.
- Applies a consistent heading/body style heuristic while copying source content.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


LETTERS = [chr(code) for code in range(ord("A"), ord("O") + 1)]
LETTERED_HEADING_RE = re.compile(r"^[a-o]\.\s+", flags=re.IGNORECASE)
WORKSTREAM_DELIVERABLE_RE = re.compile(r"^D-([A-Za-z]{3})-1(?:$|[^0-9A-Za-z])", flags=re.IGNORECASE)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Normalize drafting DOCX files and assemble section package")
    parser.add_argument("--drafting-root", required=True, help="Absolute path to drafting folder")
    parser.add_argument(
        "--workstreams-root",
        required=False,
        help="Absolute path to workstreams folder (defaults to sibling project_files/workstreams)",
    )
    parser.add_argument("--outline", required=True, help="Absolute path to outline DOCX (v3)")
    parser.add_argument("--template", required=True, help="Absolute path to Normal.dotm (presence check)")
    parser.add_argument(
        "--output-docx",
        required=True,
        help="Absolute path for assembled output DOCX",
    )
    parser.add_argument(
        "--normalized-dir",
        required=True,
        help="Absolute directory path for normalized per-source DOCX files",
    )
    return parser.parse_args()


def safe_style_name(target: DocxDocument, style_name: str | None, fallback: str = "Normal") -> str:
    if not style_name:
        return fallback
    try:
        _ = target.styles[style_name]
        return style_name
    except KeyError:
        return fallback


def infer_style_name(text: str, source_style_name: str) -> str:
    style_lower = (source_style_name or "").strip().lower()
    if style_lower.startswith("heading"):
        match = re.search(r"heading\s*(\d+)", style_lower)
        if match:
            level = max(1, min(4, int(match.group(1))))
            return f"Heading {level}"
        return "Heading 2"

    if re.match(r"^[A-O]\.\s+", text):
        return "Heading 2"
    if re.match(r"^\d+(?:\.\d+){0,2}\s+", text):
        depth = text.split(" ", 1)[0].count(".") + 1
        return f"Heading {max(1, min(4, depth + 1))}"
    if re.match(r"^[ivxlcdm]+\.\s+", text, flags=re.IGNORECASE):
        return "Heading 3"
    if len(text) <= 90 and text == text.upper() and re.search(r"[A-Z]", text):
        return "Heading 2"
    return "Normal"


def append_paragraph(target: DocxDocument, text: str, preferred_style: str, fallback_style: str = "Normal") -> None:
    style = safe_style_name(target, preferred_style, fallback=fallback_style)
    target.add_paragraph(text, style=style)


def append_doc_content(source_path: Path, target: DocxDocument) -> None:
    source = Document(str(source_path))
    for paragraph in source.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        source_style = str(paragraph.style.name or "")
        style_name = infer_style_name(text, source_style)
        append_paragraph(target, text, style_name)

    for table in source.tables:
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            row_text = " | ".join(values).strip(" |")
            if row_text:
                append_paragraph(target, row_text, "Normal")


def first_lettered_heading_paragraph(doc: DocxDocument) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if LETTERED_HEADING_RE.match(text):
            return paragraph
    return None


def trim_outline_to_section0_preface(doc: DocxDocument) -> None:
    first = first_lettered_heading_paragraph(doc)
    if first is None:
        return

    body = doc._body._element
    children = list(body.iterchildren())
    cutoff_index = None
    for idx, child in enumerate(children):
        if child is first._p:
            cutoff_index = idx
            break

    if cutoff_index is None:
        return

    for child in children[cutoff_index:]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def classify_section_files(letter_dir: Path) -> tuple[Path | None, list[Path]]:
    docx_files = sorted(letter_dir.glob("*.docx"), key=lambda path: path.name.lower())
    if not docx_files:
        return None, []

    chair_candidates = [
        path
        for path in docx_files
        if "chair" in path.stem.lower() and "preface" in path.stem.lower()
    ]
    if not chair_candidates:
        chair_candidates = [path for path in docx_files if "chair" in path.stem.lower()]

    chair = chair_candidates[0] if chair_candidates else None
    contributions = [path for path in docx_files if path != chair]

    def contribution_key(path: Path) -> tuple[int, str]:
        stem = path.stem.lower()
        is_analysis = 0 if "analysis" in stem else 1
        return (is_analysis, stem)

    contributions.sort(key=contribution_key)
    return chair, contributions


def human_title_from_filename(path: Path) -> str:
    return path.stem.replace("_", " ").replace("-", " ").strip()


def write_normalized_source_copy(source_path: Path, output_path: Path) -> None:
    normalized = Document()
    title_style = safe_style_name(normalized, "Heading 1")
    normalized.add_paragraph(human_title_from_filename(source_path), style=title_style)
    append_doc_content(source_path, normalized)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    normalized.save(str(output_path))


def add_section_divider(doc: DocxDocument, letter: str) -> None:
    doc.add_page_break()
    append_paragraph(doc, f"SECTION {letter}", "Heading 1")
    append_paragraph(doc, f"[Placeholder divider page for Section {letter}]", "Normal")
    doc.add_page_break()


def discover_workstream_deliverables(workstreams_root: Path) -> list[tuple[str, Path]]:
    found: dict[str, Path] = {}
    for docx_path in sorted(workstreams_root.rglob("*.docx"), key=lambda path: str(path).lower()):
        match = WORKSTREAM_DELIVERABLE_RE.match(docx_path.stem)
        if not match:
            continue
        code = match.group(1).upper()
        if code not in found:
            found[code] = docx_path
    return sorted(found.items(), key=lambda item: item[0])


def append_workstream_sections(
    assembled: DocxDocument,
    workstream_deliverables: list[tuple[str, Path]],
    normalized_dir: Path,
) -> int:
    if not workstream_deliverables:
        return 0

    assembled.add_page_break()
    append_paragraph(assembled, "Workstream Sections Preface", "Heading 1")
    append_paragraph(assembled, "[Preface placeholder for workstream sections]", "Normal")

    appended = 0
    for code, source_doc in workstream_deliverables:
        assembled.add_page_break()
        append_paragraph(assembled, f"WS-{code}", "Heading 1")
        append_doc_content(source_doc, assembled)
        output_doc = normalized_dir / "workstreams" / f"WS-{code}" / source_doc.name
        write_normalized_source_copy(source_doc, output_doc)
        appended += 1

    return appended


def main() -> int:
    args = parse_args()
    drafting_root = Path(args.drafting_root).resolve()
    workstreams_root = (
        Path(args.workstreams_root).resolve()
        if args.workstreams_root
        else (drafting_root.parent / "workstreams").resolve()
    )
    outline_path = Path(args.outline).resolve()
    template_path = Path(args.template).resolve()
    output_docx = Path(args.output_docx).resolve()
    normalized_dir = Path(args.normalized_dir).resolve()

    if not drafting_root.exists():
        raise SystemExit(f"drafting root not found: {drafting_root}")
    if not workstreams_root.exists():
        raise SystemExit(f"workstreams root not found: {workstreams_root}")
    if not outline_path.exists():
        raise SystemExit(f"outline docx not found: {outline_path}")
    if not template_path.exists():
        raise SystemExit(f"template not found: {template_path}")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    normalized_dir.mkdir(parents=True, exist_ok=True)

    assembled = Document(str(outline_path))
    trim_outline_to_section0_preface(assembled)

    normalized_count = 0
    included_sections = 0

    for letter in LETTERS:
        section_dir = drafting_root / f"WS-POL-{letter}"
        if not section_dir.exists() or not section_dir.is_dir():
            continue

        included_sections += 1
        add_section_divider(assembled, letter)

        chair_doc, contribution_docs = classify_section_files(section_dir)

        if chair_doc:
            append_paragraph(assembled, f"Section {letter} Chair Preface", "Heading 1")
            append_doc_content(chair_doc, assembled)
            chair_output = normalized_dir / f"WS-POL-{letter}" / chair_doc.name
            write_normalized_source_copy(chair_doc, chair_output)
            normalized_count += 1
        else:
            append_paragraph(assembled, f"Section {letter} Chair Preface", "Heading 1")
            append_paragraph(assembled, "[Missing chair preface file for this section]", "Normal")

        for contribution in contribution_docs:
            assembled.add_page_break()
            append_paragraph(assembled, human_title_from_filename(contribution), "Heading 2")
            append_doc_content(contribution, assembled)
            contribution_output = normalized_dir / f"WS-POL-{letter}" / contribution.name
            write_normalized_source_copy(contribution, contribution_output)
            normalized_count += 1

    workstream_deliverables = discover_workstream_deliverables(workstreams_root)
    appended_workstreams = append_workstream_sections(assembled, workstream_deliverables, normalized_dir)
    normalized_count += appended_workstreams

    assembled.save(str(output_docx))
    print(f"✅ Assembled drafting package: {output_docx}")
    print(f"📁 Normalized source docs: {normalized_count} files under {normalized_dir}")
    print(f"🧩 Included sections: {included_sections}/{len(LETTERS)} (A-O)")
    print(f"🗂️ Included workstreams: {appended_workstreams}")
    print(f"📝 Template path verified: {template_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
