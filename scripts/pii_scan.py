#!/usr/bin/env python3
"""Build publishable project ingest artifacts from external PROJECT_FILES_ROOT with PII gating."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import socket
import time
import urllib.error
import urllib.request
import xml.etree.ElementTree as ET
import zipfile
from datetime import date, datetime, time, timezone
from pathlib import Path
from typing import Any

try:
    from docx import Document
except ModuleNotFoundError:
    Document = None  # type: ignore[assignment]

try:
    from openpyxl import load_workbook
except ModuleNotFoundError:
    load_workbook = None  # type: ignore[assignment]

REPO_ROOT = Path(__file__).resolve().parents[1]
PUBLIC_INGEST_DIR = REPO_ROOT / "public" / "project_ingest"
PUBLIC_ARTIFACTS_DIR = PUBLIC_INGEST_DIR / "artifacts"
PUBLIC_MARKDOWN_DIR = PUBLIC_INGEST_DIR / "markdown"
PUBLIC_SPREADSHEETS_DIR = PUBLIC_INGEST_DIR / "spreadsheets"
DISCOVERY_REPORT_PATH = PUBLIC_INGEST_DIR / "discovery_report.json"
PII_REPORT_PATH = PUBLIC_INGEST_DIR / "pii_report.json"
INDEX_PATH = PUBLIC_INGEST_DIR / "index.json"
INGEST_SUMMARY_PATH = PUBLIC_INGEST_DIR / "ingest_summary.json"

PROJECT_NAME = "FCCPS AI Committee"
PIPELINE_VERSION = "2.0.0"
EXTRACTOR_VERSION = "2.0.0"

MAX_SECTION_CHARS = 1200
MAX_JSON_EXCERPT_BYTES = 20_000
MAX_XLSX_ROWS = 5000
MAX_XLSX_COLS = 100

ALLOWED_EXTENSIONS = {".md", ".txt", ".json", ".docx", ".xlsx"}
EXCLUDED_SEGMENT_EXACT = {"media"}
EXCLUDED_SEGMENT_CONTAINS = {"archive", "backup", ".git", ".venv", "__pycache__"}

EMAIL_RE = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
PHONE_RE = re.compile(r"(?<!\d)(?:\+?1[-.\s]?)?(?:\(\d{3}\)|\d{3})[-.\s]\d{3}[-.\s]\d{4}(?!\d)")
ADDRESS_RE = re.compile(
    r"\b\d{1,5}\s+[A-Za-z0-9.'-]+(?:\s+[A-Za-z0-9.'-]+){0,5}\s"
    r"(?:St|Street|Ave|Avenue|Rd|Road|Dr|Drive|Ln|Lane|Ct|Court|Blvd|Boulevard)\b",
    flags=re.IGNORECASE,
)
HANDLE_RE = re.compile(r"(?<!\w)@[A-Za-z0-9_]{2,32}\b")
STUDENT_NAME_RE = re.compile(
    r"\b(?:student|learner|pupil)\s*(?:name|initials)?\s*[:=-]\s*([A-Z][a-z]{1,20}\s+[A-Z][a-z]{1,20})\b"
)

REDACTION_TOKEN = "[REDACTED]"
TRANSFORM_VERSION = "2.0.0"
DOCX_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
DEFAULT_PROMOTION_MODEL = "gpt-4o-mini"
MAX_PROMOTION_SECTIONS = 24
MAX_PROMOTION_SECTION_CHARS = 1200
DEFAULT_OPENAI_TIMEOUT_SEC = 120
DEFAULT_OPENAI_MAX_RETRIES = 2


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def sha256_file(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def stable_slug(value: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", value.lower()).strip("_")
    return slug or "artifact"


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json_if_changed(path: Path, payload: Any) -> bool:
    serialized = json.dumps(payload, indent=2, ensure_ascii=False, sort_keys=True) + "\n"
    if path.exists() and path.read_text(encoding="utf-8") == serialized:
        return False
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(serialized, encoding="utf-8")
    return True


def write_text_if_changed(path: Path, text: str) -> bool:
    normalized = text if text.endswith("\n") else text + "\n"
    if path.exists() and path.read_text(encoding="utf-8") == normalized:
        return False
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(normalized, encoding="utf-8")
    return True


def file_mtime_utc(path: Path) -> str:
    return datetime.fromtimestamp(path.stat().st_mtime, tz=timezone.utc).isoformat().replace("+00:00", "Z")


def detect_spans(text: str) -> list[tuple[int, int, str, str]]:
    spans: list[tuple[int, int, str, str]] = []

    for match in EMAIL_RE.finditer(text):
        spans.append((match.start(), match.end(), "email", match.group(0)))
    for match in PHONE_RE.finditer(text):
        spans.append((match.start(), match.end(), "phone", match.group(0)))
    for match in ADDRESS_RE.finditer(text):
        spans.append((match.start(), match.end(), "address", match.group(0)))
    for match in HANDLE_RE.finditer(text):
        token = match.group(0)
        if "." in token:
            continue
        spans.append((match.start(), match.end(), "handle", token))
    for match in STUDENT_NAME_RE.finditer(text):
        name = match.group(1)
        if name:
            name_start = match.start(1)
            name_end = match.end(1)
            spans.append((name_start, name_end, "student_name", name))

    spans.sort(key=lambda item: (item[0], -(item[1] - item[0])))
    merged: list[tuple[int, int, str, str]] = []
    for start, end, pii_type, _raw in spans:
        if merged and start < merged[-1][1]:
            continue
        merged.append((start, end, pii_type, _raw))
    return merged


def redact_text(
    text: str,
    artifact_id: str,
    source_rel_path: str,
    section_id: str,
    field_path: str,
    findings: list[dict[str, Any]],
) -> str:
    spans = detect_spans(text)
    if not spans:
        return text

    chunks: list[str] = []
    cursor = 0
    for start, end, pii_type, _raw in spans:
        chunks.append(text[cursor:start])
        chunks.append(REDACTION_TOKEN)
        cursor = end
    chunks.append(text[cursor:])
    redacted = "".join(chunks)

    snippet = redacted
    if len(snippet) > 160:
        snippet = snippet[:157] + "..."

    for _start, _end, pii_type, _raw in spans:
        findings.append(
            {
                "artifact_id": artifact_id,
                "source_rel_path": source_rel_path,
                "section_id": section_id,
                "field_path": field_path,
                "match_type": pii_type,
                "redacted_snippet": snippet,
            }
        )

    return redacted


def chunk_text(text: str, max_chars: int = MAX_SECTION_CHARS) -> list[str]:
    normalized = re.sub(r"\r\n?", "\n", text).strip()
    if not normalized:
        return []

    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", normalized) if part.strip()]
    if not paragraphs:
        paragraphs = [normalized]

    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        if not current:
            current = paragraph
            continue
        candidate = f"{current}\n\n{paragraph}"
        if len(candidate) <= max_chars:
            current = candidate
        else:
            chunks.append(current)
            current = paragraph

    if current:
        chunks.append(current)

    split_chunks: list[str] = []
    for chunk in chunks:
        if len(chunk) <= max_chars:
            split_chunks.append(chunk)
            continue
        start = 0
        while start < len(chunk):
            split_chunks.append(chunk[start : start + max_chars])
            start += max_chars
    return split_chunks


def new_section(section_id: str, heading_path: list[str], text: str) -> dict[str, Any]:
    return {"section_id": section_id, "heading_path": heading_path, "text": text}


def extract_md_or_txt_sections(path: Path, doc_type: str) -> list[dict[str, Any]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    if doc_type == "md":
        sections: list[dict[str, Any]] = []
        heading_path: list[str] = []
        heading_stack: list[tuple[int, str]] = []
        buffer: list[str] = []

        def flush_buffer() -> None:
            if not buffer:
                return
            payload = "\n".join(buffer).strip()
            buffer.clear()
            if not payload:
                return
            for idx, chunk in enumerate(chunk_text(payload), start=1):
                section_id = f"sec_{len(sections) + 1:04d}"
                suffix = f".{idx}" if len(chunk_text(payload)) > 1 else ""
                sections.append(new_section(f"{section_id}{suffix}", list(heading_path), chunk))

        for line in re.sub(r"\r\n?", "\n", text).split("\n"):
            heading_match = re.match(r"^(#{1,6})\s+(.*\S)\s*$", line)
            if heading_match:
                flush_buffer()
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2).strip()
                while heading_stack and heading_stack[-1][0] >= level:
                    heading_stack.pop()
                heading_stack.append((level, heading_text))
                heading_path = [item[1] for item in heading_stack]
                continue
            buffer.append(line)
        flush_buffer()
        return sections

    chunks = chunk_text(text)
    return [new_section(f"sec_{idx:04d}", [], chunk) for idx, chunk in enumerate(chunks, start=1)]


def convert_excel_value(value: Any) -> Any:
    if value is None:
        return None
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, time):
        return value.isoformat()
    return str(value)


def _trim_trailing_empty_rows_and_cols(rows: list[list[Any]]) -> list[list[Any]]:
    while rows and all(cell in (None, "") for cell in rows[-1]):
        rows.pop()
    if not rows:
        return []

    max_col = 0
    for row in rows:
        for idx, cell in enumerate(row):
            if cell not in (None, ""):
                max_col = max(max_col, idx + 1)
    if max_col == 0:
        return []

    return [list(row[:max_col]) for row in rows]


def _detect_header_row(rows: list[list[Any]]) -> tuple[list[str] | None, list[list[Any]]]:
    if not rows:
        return None, rows

    first = rows[0]
    non_empty = [cell for cell in first if cell not in (None, "")]
    string_like = [cell for cell in non_empty if isinstance(cell, str) and cell.strip()]
    if len(string_like) >= 2 and len(string_like) == len(non_empty):
        header = ["" if cell is None else str(cell) for cell in first]
        return header, rows[1:]
    return None, rows


def extract_xlsx_payload(path: Path) -> dict[str, Any]:
    if load_workbook is None:
        raise RuntimeError("openpyxl is required for .xlsx ingest support")

    workbook = load_workbook(filename=str(path), read_only=True, data_only=False)
    sheets: list[dict[str, Any]] = []
    try:
        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            max_row = int(getattr(worksheet, "max_row", 0) or 0)
            max_col = int(getattr(worksheet, "max_column", 0) or 0)
            row_limit = min(max_row, MAX_XLSX_ROWS)
            col_limit = min(max_col, MAX_XLSX_COLS)

            rows: list[list[Any]] = []
            if row_limit > 0 and col_limit > 0:
                for row in worksheet.iter_rows(min_row=1, max_row=row_limit, min_col=1, max_col=col_limit, values_only=True):
                    rows.append([convert_excel_value(cell) for cell in row])

            rows = _trim_trailing_empty_rows_and_cols(rows)
            header, data_rows = _detect_header_row(rows)

            sheets.append(
                {
                    "sheet_name": str(sheet_name),
                    "header": header,
                    "rows": data_rows,
                    "row_count": len(data_rows),
                    "col_count": len(header) if isinstance(header, list) else (len(data_rows[0]) if data_rows else 0),
                    "truncated": bool(max_row > MAX_XLSX_ROWS or max_col > MAX_XLSX_COLS),
                    "row_count_original": max_row,
                    "col_count_original": max_col,
                    "row_limit": MAX_XLSX_ROWS,
                    "col_limit": MAX_XLSX_COLS,
                }
            )
    finally:
        workbook.close()

    return {
        "sheet_count": len(sheets),
        "sheets": sheets,
    }


def redact_xlsx_payload(
    payload: dict[str, Any],
    artifact_id: str,
    source_rel_path: str,
    findings: list[dict[str, Any]],
) -> dict[str, Any]:
    sheets = payload.get("sheets")
    if not isinstance(sheets, list):
        return payload

    redacted_sheets: list[dict[str, Any]] = []
    for sheet_idx, sheet in enumerate(sheets):
        if not isinstance(sheet, dict):
            continue
        header = sheet.get("header")
        redacted_header: list[str] | None = None
        if isinstance(header, list):
            redacted_header = []
            for col_idx, value in enumerate(header):
                redacted_header.append(
                    redact_text(
                        str(value),
                        artifact_id=artifact_id,
                        source_rel_path=source_rel_path,
                        section_id=f"sheet_{sheet_idx + 1:03d}",
                        field_path=f"$.sheets[{sheet_idx}].header[{col_idx}]",
                        findings=findings,
                    )
                )
        rows = sheet.get("rows")
        if not isinstance(rows, list):
            rows = []
        redacted_rows: list[list[Any]] = []
        for row_idx, row in enumerate(rows):
            if not isinstance(row, list):
                continue
            redacted_row: list[Any] = []
            for col_idx, value in enumerate(row):
                if isinstance(value, str):
                    redacted_value = redact_text(
                        value,
                        artifact_id=artifact_id,
                        source_rel_path=source_rel_path,
                        section_id=f"sheet_{sheet_idx + 1:03d}",
                        field_path=f"$.sheets[{sheet_idx}].rows[{row_idx}][{col_idx}]",
                        findings=findings,
                    )
                    redacted_row.append(redacted_value)
                else:
                    redacted_row.append(value)
            redacted_rows.append(redacted_row)

        redacted_sheet = dict(sheet)
        redacted_sheet["header"] = redacted_header
        redacted_sheet["rows"] = redacted_rows
        redacted_sheets.append(redacted_sheet)

    redacted_payload = dict(payload)
    redacted_payload["sheets"] = redacted_sheets
    return redacted_payload


def xlsx_payload_to_sections(payload: dict[str, Any]) -> list[dict[str, Any]]:
    sheets = payload.get("sheets")
    if not isinstance(sheets, list):
        return []

    sections: list[dict[str, Any]] = []
    for sheet_idx, sheet in enumerate(sheets, start=1):
        if not isinstance(sheet, dict):
            continue
        sheet_name = str(sheet.get("sheet_name") or f"Sheet {sheet_idx}")
        header = sheet.get("header")
        rows = sheet.get("rows")
        if not isinstance(rows, list):
            rows = []

        lines: list[str] = [f"sheet: {sheet_name}"]
        if isinstance(header, list):
            lines.append("header: " + " | ".join(str(cell) for cell in header))
        lines.append(
            "row_count_original: "
            + str(int(sheet.get("row_count_original") or 0))
            + ", col_count_original: "
            + str(int(sheet.get("col_count_original") or 0))
            + ", truncated: "
            + ("true" if bool(sheet.get("truncated")) else "false")
        )
        for row in rows:
            if isinstance(row, list):
                lines.append("| " + " | ".join(str(cell) for cell in row) + " |")

        text = "\n".join(lines).strip()
        for chunk_idx, chunk in enumerate(chunk_text(text), start=1):
            base = f"sec_{len(sections) + 1:04d}"
            section_id = base if chunk_idx == 1 else f"{base}.{chunk_idx}"
            sections.append(new_section(section_id, ["xlsx", sheet_name], chunk))

    return sections


def sections_to_markdown(sections: list[dict[str, Any]]) -> str:
    lines: list[str] = []
    prev_heading_path: list[str] = []

    for section in sections:
        heading_path = [str(item) for item in section.get("heading_path") or [] if str(item).strip()]
        common = 0
        for a, b in zip(prev_heading_path, heading_path):
            if a == b:
                common += 1
            else:
                break

        for idx in range(common, len(heading_path)):
            lines.append(f"{'#' * (idx + 1)} {heading_path[idx]}")
            lines.append("")

        text = str(section.get("text") or "").strip()
        if text:
            lines.append(text)
            lines.append("")

        prev_heading_path = heading_path

    return "\n".join(lines).strip() + "\n"


def extract_json_sections(path: Path) -> list[dict[str, Any]]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    try:
        parsed = json.loads(raw)
        pretty = json.dumps(parsed, indent=2, ensure_ascii=False, sort_keys=True)
    except json.JSONDecodeError:
        parsed = None
        pretty = raw

    top_level_keys: list[str] = []
    if isinstance(parsed, dict):
        top_level_keys = sorted(parsed.keys())

    if len(pretty.encode("utf-8")) > MAX_JSON_EXCERPT_BYTES:
        encoded = pretty.encode("utf-8")[:MAX_JSON_EXCERPT_BYTES]
        pretty = encoded.decode("utf-8", errors="ignore") + "\n... [TRUNCATED]"

    preamble = "json_excerpt\n"
    if top_level_keys:
        preamble += f"top_level_keys: {', '.join(top_level_keys)}\n\n"
    excerpt_text = preamble + pretty
    chunks = chunk_text(excerpt_text)
    return [new_section(f"sec_{idx:04d}", ["json_excerpt"], chunk) for idx, chunk in enumerate(chunks, start=1)]


def table_to_text(table: Any) -> str:
    lines: list[str] = []
    rendered_rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        if cells:
            rendered_rows.append(cells)
    if not rendered_rows:
        return ""

    header = rendered_rows[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row in rendered_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines).strip()


def extract_docx_sections(path: Path) -> list[dict[str, Any]]:
    if Document is None:
        return extract_docx_sections_fallback(path)

    document = Document(str(path))
    sections: list[dict[str, Any]] = []
    heading_stack: list[tuple[int, str]] = []
    current_heading_path: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = str(paragraph.style.name or "").strip().lower()
        heading_match = re.match(r"heading\s*(\d+)$", style_name)
        if heading_match:
            level = int(heading_match.group(1))
            while heading_stack and heading_stack[-1][0] >= level:
                heading_stack.pop()
            heading_stack.append((level, text))
            current_heading_path = [item[1] for item in heading_stack]
            continue

        list_prefix = ""
        if "list bullet" in style_name:
            list_prefix = "- "
        elif "list number" in style_name:
            list_prefix = "1. "

        for chunk in chunk_text(f"{list_prefix}{text}"):
            section_id = f"sec_{len(sections) + 1:04d}"
            sections.append(new_section(section_id, list(current_heading_path), chunk))

    for table in document.tables:
        rendered = table_to_text(table)
        if not rendered:
            continue
        for chunk in chunk_text(rendered):
            section_id = f"sec_{len(sections) + 1:04d}"
            sections.append(new_section(section_id, list(current_heading_path), chunk))

    return sections


def _iter_docx_text_blocks(path: Path) -> list[tuple[str, int | None]]:
    blocks: list[tuple[str, int | None]] = []
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    with zipfile.ZipFile(path) as archive:
        xml_bytes = archive.read("word/document.xml")
    root = ET.fromstring(xml_bytes)

    for paragraph in root.findall(".//w:p", ns):
        text_parts = [node.text or "" for node in paragraph.findall(".//w:t", ns)]
        text = "".join(text_parts).strip()
        if not text:
            continue
        heading_level: int | None = None
        p_style = paragraph.find(".//w:pPr/w:pStyle", ns)
        if p_style is not None:
            style_val = str(p_style.attrib.get(f"{{{ns['w']}}}val", ""))
            match = re.match(r"Heading(\d+)", style_val, flags=re.IGNORECASE)
            if match:
                heading_level = int(match.group(1))
        blocks.append((text, heading_level))

    for table in root.findall(".//w:tbl", ns):
        rows: list[str] = []
        for row in table.findall(".//w:tr", ns):
            cells: list[str] = []
            for cell in row.findall(".//w:tc", ns):
                cell_text = "".join((node.text or "") for node in cell.findall(".//w:t", ns)).strip()
                cells.append(cell_text.replace("\n", " "))
            rows.append("| " + " | ".join(cells) + " |")
        rendered = "\n".join(rows).strip()
        if rendered:
            blocks.append((rendered, None))
    return blocks


def extract_docx_sections_fallback(path: Path) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    heading_stack: list[tuple[int, str]] = []

    for text, heading_level in _iter_docx_text_blocks(path):
        if heading_level is not None:
            while heading_stack and heading_stack[-1][0] >= heading_level:
                heading_stack.pop()
            heading_stack.append((heading_level, text))
            continue

        heading_path = [item[1] for item in heading_stack]
        for chunk in chunk_text(text):
            section_id = f"sec_{len(sections) + 1:04d}"
            sections.append(new_section(section_id, heading_path, chunk))

    return sections


def extract_title(path: Path, sections: list[dict[str, Any]]) -> str:
    for section in sections:
        heading_path = section.get("heading_path")
        if isinstance(heading_path, list) and heading_path:
            first = heading_path[0]
            if isinstance(first, str) and first.strip():
                return first.strip()
    return path.stem.replace("_", " ").replace("-", " ").strip() or path.stem


def source_rel_path(path: Path, project_files_root: Path) -> str:
    return path.relative_to(project_files_root).as_posix()


def artifact_id_for(rel_path: str) -> str:
    return hashlib.sha256(rel_path.encode("utf-8")).hexdigest()[:24]


def artifact_output_name(rel_path: str, artifact_id: str) -> str:
    return f"{artifact_output_stem(rel_path, artifact_id)}.json"


def artifact_output_stem(rel_path: str, artifact_id: str) -> str:
    base = stable_slug(Path(rel_path).stem)
    short_base = base[:60] if len(base) > 60 else base
    return f"{short_base}__{artifact_id[:12]}"


def clean_output_dir(output_dir: Path, expected: set[str], glob_pattern: str) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    for existing in sorted(output_dir.glob(glob_pattern), key=lambda item: item.name.lower()):
        if existing.name not in expected:
            existing.unlink()


def discover_candidate_paths(project_files_root: Path) -> list[Path]:
    candidates: list[Path] = []
    for path in project_files_root.rglob("*"):
        if not path.is_file():
            continue
        rel_parts = [part.lower() for part in path.relative_to(project_files_root).parts]
        if any(part in EXCLUDED_SEGMENT_EXACT for part in rel_parts):
            continue
        if any(any(token in part for token in EXCLUDED_SEGMENT_CONTAINS) for part in rel_parts):
            continue
        candidates.append(path)
    candidates.sort(key=lambda item: source_rel_path(item, project_files_root).lower())
    return candidates


def should_include(path: Path) -> tuple[bool, str]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return False, "skipped_unsupported:pdf_not_ingested"
    if ext == ".pptx":
        return False, "skipped_unsupported:pptx_not_ingested"
    if ext not in ALLOWED_EXTENSIONS:
        return False, "skipped_unsupported:extension_not_allowlisted"
    return True, ""


def category_guess(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        return "document"
    if ext == ".json":
        return "data"
    if ext == ".md":
        return "notes"
    if ext == ".txt":
        return "text"
    if ext == ".xlsx":
        return "spreadsheet"
    return "other"


def extract_sections(path: Path) -> tuple[str, list[dict[str, Any]]]:
    ext = path.suffix.lower()
    if ext == ".md":
        return "md", extract_md_or_txt_sections(path, "md")
    if ext == ".txt":
        return "txt", extract_md_or_txt_sections(path, "txt")
    if ext == ".json":
        return "json", extract_json_sections(path)
    if ext == ".docx":
        return "docx", extract_docx_sections(path)
    if ext == ".xlsx":
        return "xlsx", []
    raise ValueError(f"Unsupported extension: {ext}")


def build_version_key(source_path: str, source_hash: str) -> str:
    payload = json.dumps({"source_hash": source_hash, "source_path": source_path}, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def resolve_openai_settings(enable_llm_promote: bool) -> tuple[str, str, str]:
    if not enable_llm_promote:
        return "", "", ""

    api_key = os.environ.get("OPENAI_API_KEY", "").strip() or os.environ.get("AI_INTEGRATIONS_OPENAI_API_KEY", "").strip()
    base_url = (
        os.environ.get("OPENAI_BASE_URL", "").strip()
        or os.environ.get("AI_INTEGRATIONS_OPENAI_BASE_URL", "").strip()
        or "https://api.openai.com/v1"
    )
    model = os.environ.get("PROJECT_INGEST_PROMOTION_MODEL", "").strip() or DEFAULT_PROMOTION_MODEL
    return api_key, base_url.rstrip("/"), model


def env_int(name: str, default: int, min_value: int) -> int:
    raw = os.environ.get(name, "").strip()
    if not raw:
        return default
    try:
        parsed = int(raw)
    except ValueError:
        return default
    return max(min_value, parsed)


def extract_first_json_object(text: str) -> str:
    start = text.find("{")
    if start < 0:
        raise ValueError("No JSON object start found")

    depth = 0
    in_string = False
    escape = False
    for idx in range(start, len(text)):
        ch = text[idx]
        if in_string:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == '"':
                in_string = False
            continue

        if ch == '"':
            in_string = True
            continue
        if ch == "{":
            depth += 1
            continue
        if ch == "}":
            depth -= 1
            if depth == 0:
                return text[start : idx + 1]

    raise ValueError("No complete JSON object found")


def parse_json_object_from_content(content: str) -> dict[str, Any]:
    trimmed = content.strip()
    if trimmed.startswith("```"):
        trimmed = re.sub(r"^```(?:json)?\s*", "", trimmed, flags=re.IGNORECASE)
        trimmed = re.sub(r"\s*```$", "", trimmed)

    candidates: list[str] = [trimmed]
    try:
        candidates.append(extract_first_json_object(trimmed))
    except ValueError:
        pass

    for candidate in candidates:
        if not candidate:
            continue
        try:
            parsed = json.loads(candidate)
        except json.JSONDecodeError:
            continue
        if isinstance(parsed, dict):
            return parsed

    raise RuntimeError(f"OpenAI content was not valid JSON object: {trimmed[:500]}")


def call_openai_json(
    *,
    api_key: str,
    base_url: str,
    model: str,
    system_prompt: str,
    user_prompt: str,
    max_tokens: int = 2200,
) -> dict[str, Any]:
    timeout_sec = env_int("PROJECT_INGEST_OPENAI_TIMEOUT_SEC", DEFAULT_OPENAI_TIMEOUT_SEC, 5)
    max_retries = env_int("PROJECT_INGEST_OPENAI_MAX_RETRIES", DEFAULT_OPENAI_MAX_RETRIES, 0)

    payload = {
        "model": model,
        "temperature": 0,
        "response_format": {"type": "json_object"},
        "max_tokens": max_tokens,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    }
    request = urllib.request.Request(
        f"{base_url}/chat/completions",
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )

    last_error: Exception | None = None
    total_attempts = max_retries + 1
    for attempt in range(1, total_attempts + 1):
        try:
            with urllib.request.urlopen(request, timeout=timeout_sec) as response:
                body = json.loads(response.read().decode("utf-8"))
        except urllib.error.HTTPError as exc:
            detail = exc.read().decode("utf-8", errors="ignore")
            should_retry = exc.code in {408, 409, 429, 500, 502, 503, 504}
            last_error = RuntimeError(f"OpenAI API error ({exc.code}): {detail}")
            if should_retry and attempt < total_attempts:
                sleep_sec = min(8, attempt * 2)
                print(f"   ⚠️ OpenAI HTTP {exc.code}; retrying in {sleep_sec}s ({attempt}/{total_attempts})")
                time.sleep(sleep_sec)
                continue
            raise last_error from exc
        except (urllib.error.URLError, TimeoutError, socket.timeout) as exc:
            last_error = RuntimeError(f"OpenAI request failed: {exc}")
            if attempt < total_attempts:
                sleep_sec = min(8, attempt * 2)
                print(f"   ⚠️ OpenAI request failed; retrying in {sleep_sec}s ({attempt}/{total_attempts})")
                time.sleep(sleep_sec)
                continue
            raise last_error from exc

        content = body.get("choices", [{}])[0].get("message", {}).get("content", "")
        if not isinstance(content, str) or not content.strip():
            last_error = RuntimeError("OpenAI response missing message content")
        else:
            try:
                return parse_json_object_from_content(content)
            except RuntimeError as exc:
                last_error = exc

        if attempt < total_attempts:
            sleep_sec = min(8, attempt * 2)
            print(f"   ⚠️ OpenAI returned invalid JSON; retrying in {sleep_sec}s ({attempt}/{total_attempts})")
            time.sleep(sleep_sec)

    if last_error is not None:
        raise last_error
    raise RuntimeError("OpenAI request failed with unknown error")


def promote_redaction_with_llm(
    *,
    source_rel_path: str,
    doc_type: str,
    title: str,
    redacted_sections: list[dict[str, Any]],
    pii_count: int,
    api_key: str,
    base_url: str,
    model: str,
) -> dict[str, Any]:
    section_candidates: list[dict[str, str]] = []
    for section in redacted_sections[:MAX_PROMOTION_SECTIONS]:
        section_id = str(section.get("section_id") or "")
        if not section_id:
            continue
        text = str(section.get("text") or "")
        if len(text) > MAX_PROMOTION_SECTION_CHARS:
            text = text[: MAX_PROMOTION_SECTION_CHARS - 3] + "..."
        section_candidates.append({"section_id": section_id, "text": text})

    prompt_payload = {
        "source_rel_path": source_rel_path,
        "doc_type": doc_type,
        "title": title,
        "pii_findings_count": pii_count,
        "sections": section_candidates,
    }
    system_prompt = (
        "You evaluate governance working files for publication in a committee evidence repository. "
        "Your goals are: (1) preserve only policy-relevant content, (2) aggressively remove direct personal identifiers, "
        "and (3) output strict JSON only."
    )
    user_prompt = (
        "Decide whether this file should be promoted into public ingest after redaction. "
        "If it is not materially relevant to committee governance work, set promoted=false. "
        "If promoted=true, return redacted text for each section_id that appears in input.\n\n"
        "Return JSON with this exact schema:\n"
        "{\n"
        '  "promoted": true,\n'
        '  "relevance_score": 0.0,\n'
        '  "value_add": "short phrase",\n'
        '  "rationale": "one or two sentences",\n'
        '  "redacted_sections": [{"section_id":"sec_0001","text":"..."}]\n'
        "}\n\n"
        "Rules:\n"
        "- relevance_score must be between 0 and 1.\n"
        "- Never introduce names, emails, phone numbers, social handles, addresses, or student identifiers.\n"
        "- Keep section ids unchanged; do not emit ids not present in input.\n"
        "- If promoted=false, redacted_sections may be empty.\n\n"
        f"Input JSON:\n{json.dumps(prompt_payload, ensure_ascii=False)}"
    )
    payload = call_openai_json(
        api_key=api_key,
        base_url=base_url,
        model=model,
        system_prompt=system_prompt,
        user_prompt=user_prompt,
    )

    promoted = bool(payload.get("promoted"))
    relevance_raw = payload.get("relevance_score", 0)
    relevance_score = float(relevance_raw) if isinstance(relevance_raw, (int, float)) else 0.0
    relevance_score = max(0.0, min(1.0, relevance_score))
    value_add = str(payload.get("value_add") or "").strip()
    rationale = str(payload.get("rationale") or "").strip()

    existing_by_id = {str(item.get("section_id")): item for item in redacted_sections if str(item.get("section_id"))}
    rewritten_sections: list[dict[str, Any]] = []
    raw_sections = payload.get("redacted_sections")
    if isinstance(raw_sections, list):
        for raw in raw_sections:
            if not isinstance(raw, dict):
                continue
            section_id = str(raw.get("section_id") or "")
            existing = existing_by_id.get(section_id)
            if not existing:
                continue
            rewritten_sections.append(
                {
                    "section_id": section_id,
                    "heading_path": list(existing.get("heading_path") or []),
                    "text": str(raw.get("text") or "").strip(),
                }
            )

    return {
        "promoted": promoted,
        "relevance_score": round(relevance_score, 4),
        "value_add": value_add,
        "rationale": rationale,
        "redacted_sections": rewritten_sections,
        "model": model,
    }


def markdown_with_provenance(markdown_body: str, provenance: dict[str, Any]) -> str:
    provenance_json = json.dumps(provenance, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    body = markdown_body.strip()
    return f"<!-- provenance:{provenance_json} -->\n\n{body}\n"


def resolve_project_files_root() -> Path:
    env_value = os.environ.get("PROJECT_FILES_ROOT", "").strip()
    if not env_value:
        print("❌ PROJECT_FILES_ROOT is required and must point to the external project_files folder.")
        print('   PowerShell example: $env:PROJECT_FILES_ROOT="C:\\path\\to\\project_files"')
        return Path("")
    root = Path(env_value).expanduser().resolve()
    if not root.exists() or not root.is_dir():
        print(f"❌ PROJECT_FILES_ROOT does not exist or is not a directory: {root}")
        return Path("")
    try:
        root.relative_to(REPO_ROOT)
        print(f"❌ PROJECT_FILES_ROOT must be outside this repository: {root}")
        return Path("")
    except ValueError:
        pass
    return root


def publish_artifacts(
    project_files_root: Path,
    allow_pii: bool,
    llm_promote: bool,
    openai_api_key: str,
    openai_base_url: str,
    openai_model: str,
) -> int:
    generated_at = utc_now_iso()
    candidates = discover_candidate_paths(project_files_root)

    findings: list[dict[str, Any]] = []
    discovery: list[dict[str, Any]] = []
    index_entries: list[dict[str, Any]] = []
    expected_output_filenames: set[str] = set()
    expected_md_filenames: set[str] = set()
    expected_xlsx_filenames: set[str] = set()
    type_counts: dict[str, int] = {"md": 0, "txt": 0, "json": 0, "docx": 0, "xlsx": 0}
    total_sections = 0
    skipped_due_to_pii = 0
    promoted_due_to_llm = 0

    PUBLIC_INGEST_DIR.mkdir(parents=True, exist_ok=True)

    total_candidates = len(candidates)

    for idx, path in enumerate(candidates, start=1):
        rel_path = source_rel_path(path, project_files_root)
        print(f"➡️ [{idx}/{total_candidates}] {rel_path}")
        ext = path.suffix.lower()
        include, skip_reason = should_include(path)

        entry: dict[str, Any] = {
            "source_rel_path": rel_path,
            "ext": ext,
            "size_bytes": int(path.stat().st_size),
            "mtime_utc": file_mtime_utc(path),
            "category_guess": category_guess(path),
        }

        if not include:
            entry["decision"] = "skipped"
            entry["reason_if_skipped"] = skip_reason
            discovery.append(entry)
            continue

        doc_type, extracted_sections = extract_sections(path)
        xlsx_payload: dict[str, Any] | None = None
        if doc_type == "xlsx":
            xlsx_payload = extract_xlsx_payload(path)
        rel_hash = sha256_file(path)
        artifact_id = artifact_id_for(rel_path)
        output_name = artifact_output_name(rel_path, artifact_id)
        output_stem = artifact_output_stem(rel_path, artifact_id)
        output_path = PUBLIC_ARTIFACTS_DIR / output_name
        output_rel = output_path.relative_to(REPO_ROOT).as_posix()

        redacted_sections: list[dict[str, Any]] = []
        before_count = len(findings)
        xlsx_sidecar_payload: dict[str, Any] | None = None
        if doc_type == "xlsx" and xlsx_payload is not None:
            xlsx_sidecar_payload = redact_xlsx_payload(
                xlsx_payload,
                artifact_id=artifact_id,
                source_rel_path=rel_path,
                findings=findings,
            )
            redacted_sections = xlsx_payload_to_sections(xlsx_sidecar_payload)
        else:
            for section in extracted_sections:
                section_id = str(section.get("section_id"))
                section_text = str(section.get("text") or "")
                redacted_text = redact_text(
                    section_text,
                    artifact_id=artifact_id,
                    source_rel_path=rel_path,
                    section_id=section_id,
                    field_path="$.sections[].text",
                    findings=findings,
                )
                heading_path = list(section.get("heading_path") or [])
                bounded_chunks = chunk_text(redacted_text, MAX_SECTION_CHARS) or [""]
                if len(bounded_chunks) == 1:
                    redacted_sections.append(
                        {
                            "section_id": section_id,
                            "heading_path": heading_path,
                            "text": bounded_chunks[0],
                        }
                    )
                else:
                    for chunk_idx, chunk in enumerate(bounded_chunks, start=1):
                        chunk_section_id = section_id if chunk_idx == 1 else f"{section_id}.{chunk_idx}"
                        redacted_sections.append(
                            {
                                "section_id": chunk_section_id,
                                "heading_path": heading_path,
                                "text": chunk,
                            }
                        )

        artifact_payload = {
            "artifact_id": artifact_id,
            "source_rel_path": rel_path,
            "source_hash_sha256": rel_hash,
            "file_mtime_utc": file_mtime_utc(path),
            "extracted_at_utc": generated_at,
            "doc_type": doc_type,
            "title": extract_title(path, redacted_sections),
            "sections": redacted_sections,
            "provenance": {
                "project": PROJECT_NAME,
                "pipeline_version": PIPELINE_VERSION,
                "extractor_version": EXTRACTOR_VERSION,
            },
        }

        pii_count = len(findings) - before_count
        sections_count = len(redacted_sections)
        promotion_meta: dict[str, Any] | None = None

        if pii_count > 0 and not allow_pii:
            if not llm_promote:
                skipped_due_to_pii += 1
                entry["decision"] = "skipped"
                entry["reason_if_skipped"] = "skipped_pii_findings"
                discovery.append(entry)
                continue
            try:
                print(f"   🤖 LLM promotion check: {rel_path}")
                promotion_meta = promote_redaction_with_llm(
                    source_rel_path=rel_path,
                    doc_type=doc_type,
                    title=artifact_payload["title"],
                    redacted_sections=redacted_sections,
                    pii_count=pii_count,
                    api_key=openai_api_key,
                    base_url=openai_base_url,
                    model=openai_model,
                )
            except Exception as exc:
                skipped_due_to_pii += 1
                entry["decision"] = "skipped"
                entry["reason_if_skipped"] = "skipped_pii_findings:llm_error"
                entry["promotion_error"] = str(exc)
                discovery.append(entry)
                continue

            promoted = bool(promotion_meta.get("promoted"))
            rewritten_sections = promotion_meta.get("redacted_sections")
            if promoted and isinstance(rewritten_sections, list) and rewritten_sections:
                redacted_sections = rewritten_sections
                artifact_payload["sections"] = redacted_sections
                sections_count = len(redacted_sections)
                promoted_due_to_llm += 1
            else:
                skipped_due_to_pii += 1
                entry["decision"] = "skipped"
                entry["reason_if_skipped"] = "skipped_pii_findings:not_promoted"
                discovery.append(entry)
                continue

        if promotion_meta is not None:
            artifact_payload["promotion"] = {
                "promoted": bool(promotion_meta.get("promoted")),
                "relevance_score": promotion_meta.get("relevance_score"),
                "value_add": promotion_meta.get("value_add"),
                "rationale": promotion_meta.get("rationale"),
                "model": promotion_meta.get("model"),
            }
            entry["promotion"] = {
                "promoted": bool(promotion_meta.get("promoted")),
                "relevance_score": promotion_meta.get("relevance_score"),
                "value_add": promotion_meta.get("value_add"),
            }

        associated_outputs: dict[str, str] = {}
        output_paths = [output_rel]

        if doc_type == "docx":
            md_name = f"{output_stem}.md"
            md_path = PUBLIC_MARKDOWN_DIR / md_name
            markdown_body = sections_to_markdown(redacted_sections)
            write_text_if_changed(md_path, markdown_with_provenance(markdown_body, artifact_payload["provenance"]))
            expected_md_filenames.add(md_name)
            md_rel = md_path.relative_to(REPO_ROOT).as_posix()
            associated_outputs["md_path"] = md_rel
            output_paths.append(md_rel)
            index_entries.append(
                {
                    "category": "markdown",
                    "source_path": rel_path,
                    "source_hash": rel_hash,
                    "generated_at": generated_at,
                    "version_key": build_version_key(rel_path + "::markdown", rel_hash),
                    "records_count": sections_count,
                    "pii_findings_count": pii_count,
                    "output_path": md_rel,
                }
            )

        if doc_type == "xlsx" and xlsx_sidecar_payload is not None:
            xlsx_json_name = f"{output_stem}.json"
            xlsx_json_path = PUBLIC_SPREADSHEETS_DIR / xlsx_json_name
            xlsx_sidecar_payload["provenance"] = artifact_payload["provenance"]
            write_json_if_changed(xlsx_json_path, xlsx_sidecar_payload)
            expected_xlsx_filenames.add(xlsx_json_name)
            xlsx_rel = xlsx_json_path.relative_to(REPO_ROOT).as_posix()
            associated_outputs["xlsx_json_path"] = xlsx_rel
            output_paths.append(xlsx_rel)
            records_count = 0
            for sheet in xlsx_sidecar_payload.get("sheets", []):
                if isinstance(sheet, dict):
                    rows = sheet.get("rows")
                    if isinstance(rows, list):
                        records_count += len(rows)
            index_entries.append(
                {
                    "category": "spreadsheets",
                    "source_path": rel_path,
                    "source_hash": rel_hash,
                    "generated_at": generated_at,
                    "version_key": build_version_key(rel_path + "::spreadsheets", rel_hash),
                    "records_count": records_count,
                    "pii_findings_count": pii_count,
                    "output_path": xlsx_rel,
                }
            )

        if associated_outputs:
            artifact_payload["associated_outputs"] = associated_outputs

        write_json_if_changed(output_path, artifact_payload)
        expected_output_filenames.add(output_name)

        type_counts[doc_type] += 1
        total_sections += sections_count

        entry["decision"] = "ingested"
        entry["reason_if_skipped"] = ""
        discovery.append(entry)

        index_entry: dict[str, Any] = {
            "category": "artifacts",
            "source_path": rel_path,
            "source_hash": rel_hash,
            "generated_at": generated_at,
            "version_key": build_version_key(rel_path, rel_hash),
            "records_count": sections_count,
            "pii_findings_count": pii_count,
            "output_path": output_rel,
            "output_paths": output_paths,
        }
        if associated_outputs:
            index_entry["associated_outputs"] = associated_outputs
        if promotion_meta is not None:
            index_entry["promotion"] = {
                "promoted": bool(promotion_meta.get("promoted")),
                "relevance_score": promotion_meta.get("relevance_score"),
                "value_add": promotion_meta.get("value_add"),
                "model": promotion_meta.get("model"),
            }
        index_entries.append(index_entry)

    clean_output_dir(PUBLIC_ARTIFACTS_DIR, expected_output_filenames, "*.json")
    clean_output_dir(PUBLIC_MARKDOWN_DIR, expected_md_filenames, "*.md")
    clean_output_dir(PUBLIC_SPREADSHEETS_DIR, expected_xlsx_filenames, "*.json")

    discovery_payload = {
        "generated_at": generated_at,
        "project_files_root": project_files_root.as_posix(),
        "entries": sorted(discovery, key=lambda item: str(item["source_rel_path"]).lower()),
    }
    write_json_if_changed(DISCOVERY_REPORT_PATH, discovery_payload)

    findings = sorted(
        findings,
        key=lambda item: (
            item["artifact_id"],
            item["source_rel_path"],
            item["section_id"],
            item["field_path"],
            item["match_type"],
            item["redacted_snippet"],
        ),
    )
    pii_report = {
        "allow_pii": allow_pii,
        "generated_at": generated_at,
        "total_findings": len(findings),
        "transform_version": TRANSFORM_VERSION,
        "findings": findings,
    }
    write_json_if_changed(PII_REPORT_PATH, pii_report)

    index_payload = {
        "generated_at": generated_at,
        "transform_version": TRANSFORM_VERSION,
        "entries": sorted(
            index_entries,
            key=lambda item: (
                str(item.get("category", "")).lower(),
                str(item.get("source_path", "")).lower(),
                str(item.get("output_path", "")).lower(),
            ),
        ),
    }
    write_json_if_changed(INDEX_PATH, index_payload)

    promoted_entries: list[dict[str, Any]] = []
    converted_entries: list[dict[str, Any]] = []
    for index_entry in index_payload["entries"]:
        if not isinstance(index_entry, dict):
            continue

        promotion = index_entry.get("promotion")
        if isinstance(promotion, dict) and bool(promotion.get("promoted")):
            promoted_entries.append(
                {
                    "source_path": str(index_entry.get("source_path") or ""),
                    "output_path": str(index_entry.get("output_path") or ""),
                    "pii_findings_count": int(index_entry.get("pii_findings_count") or 0),
                    "relevance_score": promotion.get("relevance_score"),
                    "value_add": str(promotion.get("value_add") or ""),
                    "model": str(promotion.get("model") or ""),
                }
            )

        associated_outputs = index_entry.get("associated_outputs")
        if not isinstance(associated_outputs, dict):
            continue

        if isinstance(associated_outputs.get("md_path"), str):
            converted_entries.append(
                {
                    "conversion_type": "docx_to_markdown",
                    "source_path": str(index_entry.get("source_path") or ""),
                    "output_path": str(associated_outputs.get("md_path") or ""),
                }
            )
        if isinstance(associated_outputs.get("xlsx_json_path"), str):
            converted_entries.append(
                {
                    "conversion_type": "xlsx_to_json",
                    "source_path": str(index_entry.get("source_path") or ""),
                    "output_path": str(associated_outputs.get("xlsx_json_path") or ""),
                }
            )

    promoted_entries.sort(key=lambda item: str(item.get("source_path", "")).lower())
    converted_entries.sort(key=lambda item: (str(item.get("conversion_type", "")).lower(), str(item.get("source_path", "")).lower()))

    ingest_summary_payload = {
        "generated_at": generated_at,
        "transform_version": TRANSFORM_VERSION,
        "counts": {
            "promoted": len(promoted_entries),
            "converted": len(converted_entries),
        },
        "promoted": promoted_entries,
        "converted": converted_entries,
    }
    write_json_if_changed(INGEST_SUMMARY_PATH, ingest_summary_payload)

    print("✅ Project ingest artifact publish complete")
    print(f"📄 Candidates scanned: {len(candidates)}")
    print(
        "🧾 Ingested by type: "
        + f"md={type_counts['md']}, txt={type_counts['txt']}, json={type_counts['json']}, "
        + f"docx={type_counts['docx']}, xlsx={type_counts['xlsx']}"
    )
    print(f"🧱 Total sections: {total_sections}")
    print(f"🕵️ PII findings: {len(findings)}")
    print(f"⏭️ Skipped due to PII: {skipped_due_to_pii}")
    if llm_promote:
        print(f"🚀 Promoted via LLM redaction: {promoted_due_to_llm}")
    return 0


def main() -> None:
    parser = argparse.ArgumentParser(description="Publish scrubbed project ingest artifacts with PII gating")
    parser.add_argument("--allow-pii", action="store_true", help="Allow publish step to succeed even when PII findings are detected")
    parser.add_argument(
        "--llm-promote",
        action="store_true",
        help="For PII-flagged files, call OpenAI to decide relevance and produce enhanced redaction for promotion",
    )
    args = parser.parse_args()

    project_files_root = resolve_project_files_root()
    if not project_files_root:
        raise SystemExit(1)

    openai_api_key, openai_base_url, openai_model = resolve_openai_settings(args.llm_promote)
    if args.llm_promote and not openai_api_key:
        print("❌ --llm-promote requires OPENAI_API_KEY (or AI_INTEGRATIONS_OPENAI_API_KEY) in your environment.")
        raise SystemExit(1)

    raise SystemExit(
        publish_artifacts(
            project_files_root=project_files_root,
            allow_pii=args.allow_pii,
            llm_promote=args.llm_promote,
            openai_api_key=openai_api_key,
            openai_base_url=openai_base_url,
            openai_model=openai_model,
        )
    )


if __name__ == "__main__":
    main()
